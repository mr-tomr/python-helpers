# Use to both list and count the names of findings in a pentest report which uses this header structure
# Also count the number of each occurence for the summary table

import sys
import os
from docx import Document

def extract_findings_and_instances(docx_path, output_txt):
    document = Document(docx_path)
    headings = []

    for para in document.paragraphs:
        style = para.style.name
        text = para.text.strip()
        if style == 'Heading 3' and text:
            headings.append('\n' + text)
        elif style == 'Heading 4' and text:
            headings.append(text)

    with open(output_txt, 'w', encoding='utf-8') as f:
        for line in headings:
            f.write(line + '\n')

    print(f"[+] Extracted Findings and Instances to '{output_txt}'")

def extract_finding_and_count(docx_path, output_txt):
    document = Document(docx_path)
    lines = []
    paragraphs = document.paragraphs
    i = 0
    total = len(paragraphs)

    while i < total:
        para = paragraphs[i]
        style = para.style.name
        text = para.text.strip()

        if style == 'Heading 3' and text:
            h3_text = text
            h4_count = 0
            i += 1

            while i < total:
                next_para = paragraphs[i]
                next_style = next_para.style.name
                next_text = next_para.text.strip()

                if next_style == 'Heading 4' and next_text:
                    h4_count += 1
                    i += 1
                elif next_style == 'Heading 3':
                    break
                else:
                    i += 1

            lines.append(f"{h3_text}\t{h4_count}")
        else:
            i += 1

    with open(output_txt, 'w', encoding='utf-8') as f:
        for line in lines:
            f.write(line + '\n')

    print(f"[+] Extracted {len(lines)} Finding names and count to '{output_txt}'")

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python merged_extract.py <input.docx>")
        sys.exit(1)

    input_file = sys.argv[1]
    if not os.path.isfile(input_file):
        print(f"[!] File not found: {input_file}")
        sys.exit(1)

    base = os.path.splitext(os.path.basename(input_file))[0]
    findings_and_instances_file = f"{base}_findings_and_instances.txt"
    finding_and_count_file = f"{base}_finding_and_count.txt"

    extract_findings_and_instances(input_file, findings_and_instances_file)
    extract_finding_and_count(input_file, finding_and_count_file)
